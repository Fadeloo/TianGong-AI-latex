import requests
import json
import os
import glob
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from nougat_latex_processor import (
    process_images_to_latex,
)  # Import your nougat_latex_processor module

# Load environment variables from .env file
load_dotenv()

# Get API credentials and directory paths from environment variables
TEXTIN_API_ID = os.getenv("TEXTIN_API_ID")
TEXTIN_API_CODE = os.getenv("TEXTIN_API_CODE")
INPUT_DIRECTORY = os.getenv("INPUT_DIRECTORY")
OUTPUT_DIRECTORY = os.getenv("OUTPUT_DIRECTORY")


# Read file content
def get_file_content(filePath):
    with open(filePath, "rb") as fp:
        return fp.read()


# PDF parsing
class CommonOcr(object):
    def __init__(self, img_path):
        self._app_id = TEXTIN_API_ID
        self._secret_code = TEXTIN_API_CODE
        self._img_path = img_path

    def recognize(self):
        url = "https://api.textin.com/ai/service/v1/pdf_to_markdown?get_image=objects"
        head = {"x-ti-app-id": self._app_id, "x-ti-secret-code": self._secret_code}
        try:
            image = get_file_content(self._img_path)
            result = requests.post(url, data=image, headers=head)
            return result
        except Exception as e:
            return e


# Download image and save to local
def download_and_save_image(image_url, save_dir):
    response = requests.get(image_url)
    if response.status_code == 200:
        file_path = os.path.join(save_dir, os.path.basename(image_url))
        with open(file_path, "wb") as file:
            file.write(response.content)
        return file_path
    else:
        print("Unable to download image")
        return None


# Delete local file
def delete_file(file_path):
    if os.path.exists(file_path):
        os.remove(file_path)


# Create or get custom table and paragraph styles
def create_custom_styles(doc):
    styles = doc.styles
    if "CustomTableStyle" not in styles:
        table_style = styles.add_style("CustomTableStyle", WD_STYLE_TYPE.TABLE)
        table_style.font.name = "宋体"
        table_style.font.size = Pt(10)
        table_style.paragraph_format.space_before = Pt(0)
        table_style.paragraph_format.space_after = Pt(0)
        table_style.paragraph_format.line_spacing = 1.5

    if "CustomTable" not in styles:
        paragraph_style = styles.add_style(
            "CustomTable", WD_STYLE_TYPE.PARAGRAPH
        )
        paragraph_style.font.name = "宋体"
        paragraph_style.font.size = Pt(10)
        paragraph_style.paragraph_format.space_before = Pt(0)
        paragraph_style.paragraph_format.space_after = Pt(0)
        paragraph_style.paragraph_format.line_spacing = 1.5

    return styles["CustomTableStyle"], styles["CustomTable"]


# Generate table in docx
def html_table_to_docx(html_content, doc):
    # Parse table
    soup = BeautifulSoup(html_content, "html.parser")

    # Find tables in HTML
    tables = soup.find_all("table")
    for table in tables:
        rows = table.find_all("tr")
        # Get column count
        column_num = 0
        for cell in rows[0].find_all(["td", "th"]):
            column_num += int(cell.get("colspan", 1))
        word_table = doc.add_table(rows=len(rows), cols=column_num)

        # Set table and paragraph styles explicitly
        table_style, paragraph_style = create_custom_styles(doc)
        word_table.style = table_style

        # Apply paragraph style to each cell
        for row in word_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = paragraph_style

        # Validate table style application
        print(f"Applied table style: {word_table.style.name}")

        # Record merged cell positions and spans
        merged_cells = []
        # Record merged cell start and end
        merged_list = []

        # Create table first
        # Iterate over each row
        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(["td", "th"])
            word_row = word_table.rows[row_idx]
            cell_idx = 0

            # Iterate over each column
            for cell in cells:
                colspan = int(cell.get("colspan", 1))
                rowspan = int(cell.get("rowspan", 1))

                # Skip merged cells (rows)
                while (row_idx, cell_idx) in merged_cells:
                    cell_idx += 1

                # Write cell content first
                word_cell = word_row.cells[cell_idx]
                word_cell.text = cell.get_text(strip=True)

                # Apply paragraph style to cell paragraphs
                for paragraph in word_cell.paragraphs:
                    paragraph.style = paragraph_style

                # Record merged cells
                if colspan > 1 or rowspan > 1:
                    for r in range(rowspan):
                        for c in range(colspan):
                            merged_cells.append((row_idx + r, cell_idx + c))
                    merged_list.append(
                        [
                            (row_idx, cell_idx),
                            (row_idx + rowspan - 1, cell_idx + colspan - 1),
                        ]
                    )

                # Skip merged cells (columns)
                cell_idx += colspan

        # Merge cells
        for merged_pairs in merged_list:
            cell_1st = word_table.rows[merged_pairs[0][0]].cells[merged_pairs[0][1]]
            cell_2nd = word_table.rows[merged_pairs[1][0]].cells[merged_pairs[1][1]]
            cell_1st.merge(cell_2nd)


# Create or get "No Spacing" style
def get_no_spacing_style(doc):
    styles = doc.styles
    if "No Spacing" not in styles:
        style = styles.add_style("No Spacing", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = Pt(1)
    return styles["No Spacing"]


# Generate docx document
def get_title_level(body_text):
    title_level = []
    if len(body_text) < 30:
        number_list = body_text.split(" ")[0].split(".")
        if len(number_list) < 4:
            try:
                for i in number_list:
                    title_level.append(int(i))
            except:
                title_level = []
    return title_level


def docs_output(doc, list_name, temp_dir):
    is_main_body = 0
    title_level = []
    no_spacing_style = get_no_spacing_style(doc)  # Get "No Spacing" style
    for i in range(len(list_name)):
        # Text
        if list_name[i]["type"] == "paragraph":
            if is_main_body == 0:
                # Footer
                if list_name[i]["content"] == 1:
                    # Ignore footers shorter than 3 characters
                    if len(list_name[i]["text"]) > 2:
                        doc.add_paragraph(list_name[i]["text"], style=no_spacing_style)
                # Heading
                elif list_name[i]["outline_level"] >= 0:
                    # Determine if it is the main text
                    if list_name[i]["text"][0] in "0123456789":
                        is_main_body = 1
                        title_level = get_title_level(list_name[i]["text"])
                        doc.add_heading(list_name[i]["text"], level=len(title_level))
                        continue
                    else:
                        doc.add_heading(list_name[i]["text"], level=1)
                # Main text
                else:
                    doc.add_paragraph(list_name[i]["text"], style=no_spacing_style)
            if is_main_body == 1:
                if list_name[i]["content"] == 1:
                    if len(list_name[i]["text"]) > 2:
                        doc.add_paragraph(list_name[i]["text"], style=no_spacing_style)
                else:
                    title_level = get_title_level(list_name[i]["text"])
                    if title_level == []:
                        doc.add_paragraph(list_name[i]["text"], style=no_spacing_style)
                    else:
                        doc.add_heading(list_name[i]["text"], level=len(title_level))
        # Image
        elif list_name[i]["type"] == "image":
            image_url = list_name[i]["image_url"]
            image_path = download_and_save_image(
                image_url, temp_dir
            )  # Download image and save to temporary directory
            if image_path:
                latex_results = process_images_to_latex(
                    img_dir=temp_dir
                )  # Process downloaded images
                for latex in latex_results:
                    doc.add_paragraph(
                        latex, style=no_spacing_style
                    )  # Add LaTeX code to document
                delete_file(image_path)  # Delete downloaded image file
        # Table
        elif list_name[i]["type"] == "table":
            print("Writing table")
            html_table_str = list_name[i]["text"]
            html_table_to_docx(html_table_str, doc)
        else:
            print("New type found: " + list_name[i]["type"])


# Process a single file
def process_single_file(file_path, output_directory):
    with tempfile.TemporaryDirectory() as temp_dir:
        response = CommonOcr(file_path)
        try:
            pdf_result = response.recognize()
            data_dict = json.loads(pdf_result.text)
            data_list = data_dict["result"]["detail"]
            print(f"{os.path.basename(file_path)} parsing completed")
        except Exception as e:
            print(f"{os.path.basename(file_path)} parsing failed")
            print(e)
            return

        try:
            doc = Document()
            docs_output(doc, data_list, temp_dir)
            output_file_path = os.path.join(
                output_directory,
                f"{os.path.splitext(os.path.basename(file_path))[0]}.docx",
            )
            doc.save(output_file_path)
            print(f"{os.path.basename(file_path)} document generated successfully!")
        except Exception as e:
            print(f"{os.path.basename(file_path)} document generation failed")
            print(e)


# Process all files in the specified directory
def process_all_files_in_directory(input_directory, output_directory):
    with ThreadPoolExecutor() as executor:
        futures = []
        for file_path in glob.glob(os.path.join(input_directory, "*")):
            if os.path.isfile(file_path):
                futures.append(
                    executor.submit(process_single_file, file_path, output_directory)
                )

        for future in as_completed(futures):
            try:
                future.result()
            except Exception as e:
                print(f"Error processing file: {e}")


# Main function
if __name__ == "__main__":
    process_all_files_in_directory(INPUT_DIRECTORY, OUTPUT_DIRECTORY)

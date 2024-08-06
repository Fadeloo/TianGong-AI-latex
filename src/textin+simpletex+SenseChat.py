import requests
import json
from docx import Document
import datetime
from random import Random
import hashlib
from bs4 import BeautifulSoup


# api的账号和密码
# 合合textin
TEXTIN_API_ID = "textin api id"
TEXTIN_API_CODE = "textin api code"
# simpletex
SIMPLETEX_APP_ID = "simpletex api id"
SIMPLETEX_APP_SECRET = "simpletex api code"
# 商汤sensechat
SENSECHAT_API_ID = "sensechat api id"
SENSECHAT_API_CODE = "sensechat api code"


# 读取文件
def get_file_content(filePath):
    with open(filePath, 'rb') as fp:
        return fp.read()


# pdf解析
# 调用textin api，需要自己申请账号和密码，参考网站：https://www.textin.com/document/pdf_to_markdown
class CommonOcr(object):
    def __init__(self, img_path):
        self._app_id = TEXTIN_API_ID
        self._secret_code = TEXTIN_API_CODE
        self._img_path = img_path

    def recognize(self):
        # 通用文档解析
        url = 'https://api.textin.com/ai/service/v1/pdf_to_markdown'
        url = url + '?' + 'get_image=objects'
        head = {}
        try:
            image = get_file_content(self._img_path)
            head['x-ti-app-id'] = self._app_id
            head['x-ti-secret-code'] = self._secret_code
            result = requests.post(url, data=image, headers=head)
            return result
        except Exception as e:
            return e


# latex转换
# 调用simpletex api，参考网站：https://simpletex.cn/api_doc
def download_image(image_url):
    response = requests.get(image_url)
    if response.status_code == 200:
        return response.content
    else:
        print("无法下载图片")
        return None


def random_str(randomlength=16):
    str = ''
    chars = 'AaBbCcDdEeFfGgHhIiJjKkLlMmNnOoPpQqRrSsTtUuVvWwXxYyZz0123456789'
    length = len(chars) - 1
    random = Random()
    for i in range(randomlength):
        str += chars[random.randint(0, length)]
    return str


def get_req_data(req_data, appid, secret):
    header = {}
    header["timestamp"] = str(int(datetime.datetime.now().timestamp()))
    header["random-str"] = random_str(16)
    header["app-id"] = appid
    pre_sign_string = ""
    sorted_keys = list(req_data.keys()) + list(header)
    sorted_keys.sort()
    for key in sorted_keys:
        if pre_sign_string:
            pre_sign_string += "&"
        if key in header:
            pre_sign_string += key + "=" + str(header[key])
        else:
            pre_sign_string += key + "=" + str(req_data[key])

    pre_sign_string += "&secret=" + secret
    header["sign"] = hashlib.md5(pre_sign_string.encode()).hexdigest()
    return header, req_data


# 商汤大模型解释，使用SenseChat-Vision大模型，需要账号和密码，参考网站：https://console.sensecore.cn/help/docs/model-as-a-service/nova/vision/ChatCompletions/
sensenova.access_key_id = SENSECHAT_API_ID
sensenova.secret_access_key = SENSECHAT_API_CODE
def get_image_description(image_url, prompt_str):
    resp = sensenova.ChatCompletion.create(
        model="SenseChat-Vision",
        max_new_tokens=2048,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": image_url
                    },
                    {
                        "type": "text",
                        "text": "这是" + prompt_str + "，请你对这张图的工作原理或流程等做出详细解释，字数在500字左右"
                    }
                ]
            }
        ],
        repetition_penalty=1.0,
        temperature=0.3,
        top_p=0.7,
        stream=False,
        user="image-interpreter"
    )
    return resp


# 表格生成
def html_table_to_docx(html_content, doc):
    # 解析表格
    soup = BeautifulSoup(html_content, 'html.parser')

    # 查找HTML中的表格
    tables = soup.find_all('table')
    for table in tables:
        rows = table.find_all('tr')
        # 获取列表
        column_num = 0
        for cell in rows[0].find_all(['td', 'th']):
            column_num = column_num + int(cell.get('colspan', 1))
        word_table = doc.add_table(rows=len(rows), cols=column_num)

        # 记录合并的单元格位置和跨度
        merged_cells = []
        # 记录合并的单元格首尾
        merged_list = []

        # 先创建表格
        # 遍历每一行
        for row_idx, tr in enumerate(rows):
            cells = tr.find_all(['td', 'th'])
            word_row = word_table.rows[row_idx]
            cell_idx = 0

            # 遍历每一列
            for cell in cells:
                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))

                # 跳过合并的单元格(行)
                while (row_idx, cell_idx) in merged_cells:
                    cell_idx += 1

                # 先写入单元格的内容
                word_cell = word_row.cells[cell_idx]
                word_cell.text = cell.get_text(strip=True)

                # 记录合并的单元格
                if colspan > 1 or rowspan > 1:
                    for r in range(rowspan):
                        for c in range(colspan):
                            merged_cells.append((row_idx + r, cell_idx + c))
                    merged_list.append(
                        [(row_idx, cell_idx), (row_idx+rowspan-1, cell_idx+colspan-1)])

                # 跳过合并的单元格（列）
                cell_idx += colspan

        # 再合并单元格
        # print(merged_cells)
        # print(merged_list)
        for merged_pairs in merged_list:
            cell_1st = word_table.rows[merged_pairs[0][0]].cells[merged_pairs[0][1]]
            cell_2nd = word_table.rows[merged_pairs[1][0]].cells[merged_pairs[1][1]]
            cell_1st.merge(cell_2nd)


# 生成doc文档
def get_title_level(body_text):
    title_level = []
    if len(body_text) < 30:
        number_list = body_text.split(' ')[0].split('.')
        if len(number_list) < 4:
            try:
                for i in number_list:
                    title_level.append(int(i))
            except:
                title_level = []
    return title_level

def docs_output(doc, list_name):
    page_num = 1
    is_main_body = 0
    title_level = []
    print(f"第{page_num}页生成中")
    for i in range(len(list_name)):
        # 换页
        if list_name[i]['page_id'] > page_num:
            doc.add_page_break()
            page_num = page_num + 1
            print(f"第{page_num}页生成中")
        # 文字
        if list_name[i]['type'] == 'paragraph':
            if is_main_body == 0:
                # 页脚
                if list_name[i]['content'] == 1:
                    # 忽略长度小于3的页脚
                    if len(list_name[i]['text']) > 2:
                        doc.add_paragraph(list_name[i]['text'])
                # 标题
                elif list_name[i]['outline_level'] >= 0:
                    # 判断是否为文本主体
                    if list_name[i]['text'][0] in '0123456789':
                        is_main_body = 1
                        title_level = get_title_level(list_name[i]['text'])
                        doc.add_heading(list_name[i]['text'], level = len(title_level))
                        continue
                    else:
                        doc.add_heading(list_name[i]['text'], level = 1)
                # 正文
                else:
                    doc.add_paragraph(list_name[i]['text'])
            if is_main_body == 1:
                if list_name[i]['content'] == 1:
                    if len(list_name[i]['text']) > 2:
                        doc.add_paragraph(list_name[i]['text'])
                else:
                    title_level = get_title_level(list_name[i]['text'])
                    if title_level == []:
                        doc.add_paragraph(list_name[i]['text'])
                    else:
                        doc.add_heading(list_name[i]['text'], level = len(title_level))
        # 图片
        elif list_name[i]['type'] == 'image':
            image_url = list_name[i]['image_url']
            image_data = download_image(image_url)
            # 图片解释
            if list_name[i+1]['text'][0] == '图':
                print("图片解释中")
                try:
                    description = get_image_description(
                        image_url, list_name[i+1]['text'])
                    text = description.data.choices[0].message
                except Exception as e:
                    print(e)
            # 公式转换
            else:
                print("公式转换中")
                pdf_file = {"file": image_data}
                data = {}
                header, data = get_req_data(data, SIMPLETEX_APP_ID, SIMPLETEX_APP_SECRET)
                try:
                    res = requests.post(
                        "https://server.simpletex.cn/api/latex_ocr", files=pdf_file, data=data, headers=header)
                    equation = json.loads(res.text)['res']['latex']
                except Exception as e:
                    print(e)
                text = '$' + equation + '$'
            doc.add_paragraph(text)
        # 表格
        elif list_name[i]['type'] == 'table':
            print("表格书写中")
            html_table_str = list_name[i]['text']
            html_table_to_docx(html_table_str, doc)

        else:
            print("有新的类型：" + list_name[i]['type'])


# %% 主程序
if __name__ == "__main__":
    response = CommonOcr(r'输入的pdf文件路径')
    try:
        pdf_result = response.recognize()
        # print(pdf_result.text)
        data_dict = (json.loads(pdf_result.text))
        data_list = data_dict['result']['detail']
        print("pdf解析完成")
    except Exception as e:
        print("pdf解析失败")
        print(e)


    try:
        doc = Document()
        docs_output(doc, data_list)
        doc.save('输出的docx文件路径')
        print("文档生成成功！")
    except Exception as e:
        print("文档生成失败")
        print(e)



